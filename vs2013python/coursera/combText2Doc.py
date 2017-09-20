from docx import Document
from docx.shared import Inches
import os

# pip install python-docx ; 0.8.6
# combine text in subfolders in one doc
PATH = r'''Z:\1'''

def CombText2Doc(dirname):
    if dirname[0] == '.':
        return
    os.chdir(dirname)
    files = os.listdir()
    files.sort()

    document = Document()
    document.add_heading(dirname, 0)
    hasContent = False

    for fn in files:
        if not os.path.isfile(fn):
            continue
        paths = os.path.splitext(fn)
        if paths[1] == '.txt':
            print('proc file %s' % fn)
            hasContent = True
            document.add_heading(paths[0], level=1)
            with open(fn, 'r') as f:
                for s in f: 
                    s = s.strip()
                    if len(s) == 0:
                        continue
                    document.add_paragraph(s.strip()).paragraph_format.first_line_indent = Inches(0.25)
                    document.add_paragraph('')
    if hasContent:
        document.save(dirname + '.docx')
    
    os.chdir('..')

if __name__ == '__main__':
    os.chdir(PATH)
    files = os.listdir()
    for fn in files:
        if os.path.isdir(fn):
            print('****proc dir %s' % fn)
            CombText2Doc(fn)